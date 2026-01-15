import os
from docx import Document
from rich.console import Console
from src.services.gemini import GeminiClient

console = Console()

class DocxHandler:
    def __init__(self, gemini_client: GeminiClient, target_lang: str):
        self.client = gemini_client
        self.target_lang = target_lang

    def process(self, input_path: str):
        """
        Reads a DOCX file, translates paragraphs and table cells,
        and saves the result.
        """
        output_path = f"{os.path.splitext(input_path)[0]}_translated.docx"

        doc = Document(input_path)

        # Translate Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # To preserve runs (bold/italic) is complex in Phase 1.
                # MVP Strategy: Translate the whole paragraph text and replace it.
                # This loses inner formatting like "This is **bold** text" -> "這是粗體文字" (bold lost or applied to whole)
                # Better Phase 1: Translate text, clear runs, add one run with new text.
                # Or: Iterate runs? If we iterate runs, we fragment sentences.
                # "Hello" (Run 1) " World" (Run 2). Translating separately is bad.

                # Decision for MVP: Translate full paragraph text.
                # Attempt to keep style of the first run if possible, or just the paragraph style.

                original_text = para.text
                translated_text = self.client.translate_text(original_text, self.target_lang)

                # Clear existing content
                para.clear()
                # Add translated text
                para.add_run(translated_text)

        # Translate Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Cells can contain paragraphs too
                    for para in cell.paragraphs:
                        if para.text.strip():
                            original_text = para.text
                            translated_text = self.client.translate_text(original_text, self.target_lang)
                            para.clear()
                            para.add_run(translated_text)

        doc.save(output_path)
        return output_path
